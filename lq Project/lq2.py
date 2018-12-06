'''
蓝桥杯的爬虫
'''

from bs4 import BeautifulSoup
from urllib import request
from selenium import webdriver
import time
import re
from PIL import Image
import os
import docx
import requests
from docx import Document
from docx.shared import Inches


#
# def Get_Opener():
#     # 创建一个cookiejar对象
#     cookiejar=CookieJar()
#     # 使用cookiejar创建一个httpcookieprocessor对象
#     handle=request.HTTPCookieProcessor(cookiejar)
#     # 使用handle创建一个opener
#     opener=request.build_opener(handle)
#
#     return opener




# def Get_Url():
#     url="http://lx.lanqiao.cn/problem.page?gpid="
#     for i in range(412,413):
#         Save_Data(url+"T{0}".format(i))



def Login():
    url="http://dasai.lanqiao.cn/pages/account/login_other.html?backurl=http%3A//lx.lanqiao.cn/lqloginfinished.page%3Fredir%3D/problem.page"
    header = {
        "User - Agent": "Mozilla / 5.0(WindowsNT10.0;WOW64) AppleWebKit / 537.36(KHTML, likeGecko) Chrome / 65.0.3325.181 Safari / 537.36",
    }

    driver = webdriver.Firefox()

    driver.get(url)
    #
    #
    # handles = driver.window_handles
    # driver.switch_to_window(handles[1])

    driver.find_element_by_name('login').send_keys("18877423546")
    driver.find_element_by_name('pwd').send_keys("cjy1995309")
    time.sleep(6)
    driver.find_element_by_xpath("//*[@class='sub mt10']").click()

    # print(driver.page_source)
    time.sleep(5)
    # driver.find_element_by_xpath("//a[@href='/problemsets.page']").click()

    url1 = "http://lx.lanqiao.cn/problem.page?gpid="
    # now_handle = driver.current_window_handle
    for i in range(501,502):
        Save_Data(url1 + "T{0}".format(i),i,driver)
        time.sleep(3)


def Save_Data(url,i,driver):
    print(i)
    driver.get(url)
    document=Document()


    handles = driver.window_handles
    driver.switch_to_window(handles[0])

    html=driver.page_source

    # print(html)
    soup = BeautifulSoup(html, 'lxml')

    print("--------")
    img1 = soup.find_all("div", class_="des")
    try:
        img = str(soup.find_all('div', {"class": "pddata"})).split('src="')[1].split('"/')[0]

        img = img.split('"')[0]
        tu = 'http://lx.lanqiao.cn' + img
        print(tu)
        with open('../data/{0}.jpg'.format(i), 'wb')as f:
            response = requests.get(tu).content
            f.write(response)
            f.close()
    except Exception:
        print("NOT FOUND!")


    re_h = re.compile('</?\w+[^>]*>')  # HTML标签
    re_br= re.compile('<br\s*?/?>')  # 处理换行
    # 处理&lt
    re_lt=re.compile('&lt;')
    re_gt = re.compile('&gt;')
    re_amp = re.compile('&amp;')
    re_left = re.compile("\[")
    re_right = re.compile(']')


    s=re_br.sub("\n", str(img1))
    s = re_h.sub('', s)
    s=re_lt.sub('<',s)
    s = re_gt.sub('>', s)
    s = re_amp.sub('&', s)
    s = re_left.sub('', s)
    s = re_right.sub('', s)
    # print(s)

    try:
        document.add_paragraph(s)  # 向文档里添加文字
        document.add_picture('{0}.jpg'.format(i))
    except Exception:
        print('img not found')
    document.save('../data/{0}.doc'.format(i))  # 保存文档

    # os.remove(img_name)  # 删除保存在本地的图片


if __name__ == '__main__':
    Login()
