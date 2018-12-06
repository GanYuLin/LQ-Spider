import requests
from bs4 import BeautifulSoup
import os
import docx
from docx import Document
from docx.shared import Inches

url = 'https://www.qiushibaike.com/article/119757360'
html = requests.get(url).content
soup = BeautifulSoup(html,'html.parser')
wen = soup.find('div',{"class":"content"}).text
img = str(soup.find('div',{"class":"thumb"})).split('src="')[1].split('"/')[0]
tu = 'https:' + img
img_name = img.split('/')[-1]
print(img)
#保存图片至本地
with open(img_name,'wb')as f:
    response = requests.get(tu).content
    f.write(response)
    f.close()

document = Document()
document.add_paragraph(wen)#向文档里添加文字
document.add_picture(img_name)#向文档里添加图片
document.save('tuwen.doc')#保存文档
os.remove(img_name)#删除保存在本地的图片
